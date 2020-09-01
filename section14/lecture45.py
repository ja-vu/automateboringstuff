import docx
d = docx.Document('c:\\Workspace\\automateboringstuff\\section14\\demo.docx')
print(d.paragraphs)
p = d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[0].bold)
print(p.runs[3].italic)
p.runs[3].underline = True
p.runs[3].text = 'italic and underline'

print(p.runs[3].text)
d.save('c:\\Workspace\\automateboringstuff\\section14\\demo2.docx')
p.style = 'Title'
d.save('c:\\Workspace\\automateboringstuff\\section14\\demo3.docx')


d2 = docx.Document()
d2.add_paragraph('Hello this is a paragraph')
d2.add_paragraph('This is another paragraph')
d2.save('c:\\Workspace\\automateboringstuff\\section14\\demo4.docx')
p2 = d2.paragraphs[0]
p2.add_run('This is a new run')
p2.runs[1].bold = True
d2.save('c:\\Workspace\\automateboringstuff\\section14\\demo5.docx')

